from pathlib import Path

from docx import Document


def read_docx(file_path: str | Path) -> str:
    """
    DOCX dosyasındaki metinleri okur ve tek bir metin olarak döndürür.
    """

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"Dosya bulunamadı: {path}")

    if path.suffix.lower() != ".docx":
        raise ValueError(f"Desteklenmeyen dosya türü: {path.suffix}")

    document = Document(path)

    paragraphs: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    # Dokümandaki tabloları da okur.
    for table in document.tables:
        for row in table.rows:
            row_values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if row_values:
                paragraphs.append(" | ".join(row_values))

    return "\n".join(paragraphs)