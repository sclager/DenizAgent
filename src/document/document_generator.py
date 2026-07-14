from __future__ import annotations

import re
import unicodedata
from datetime import date, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.table import _Cell


PROJECT_ROOT = Path(__file__).resolve().parents[2]

TEMPLATE_PATH = (
    PROJECT_ROOT
    / "templates"
    / "DEVCR_XXX_TEMPLATE.docx"
)

DOCUMENTS_DIRECTORY = PROJECT_ROOT / "documents"


FIELD_LABELS = {
    "itsm_no": "ITSM NO",
    "request_date": "Talep Tarihi",
    "customer_name": "Talepte Bulunan Müşteri",
    "country": "Talep Hangi Ülkede Kullanılacak?",
    "request": "Talep",
    "request_description": "Talebin Tanımlanması",
    "detailed_description": "İsteğin ayrıntılı açıklaması",
    "cash_register_work": (
        "Kasa tarafında yapılacak çalışmalar/işlem akışı"
    ),
    "service_work": (
        "Servis tarafında yapılacak çalışmalar/işlem akışı"
    ),
    "database_information": (
        "Database içinde istenilen veriler yerleri nelerdir."
    ),
    "success_failure_scenarios": (
        "Geliştirilecek uygulama için BAŞARILI ve BAŞARISIZ "
        "işlemler için detaylı işleyiş açıklaması:"
    ),
}


TURKISH_CHARACTER_MAP = str.maketrans(
    {
        "ç": "c",
        "Ç": "C",
        "ğ": "g",
        "Ğ": "G",
        "ı": "i",
        "İ": "I",
        "ö": "o",
        "Ö": "O",
        "ş": "s",
        "Ş": "S",
        "ü": "u",
        "Ü": "U",
    }
)


INVALID_XML_CHARACTERS = re.compile(
    "["
    "\x00-\x08"
    "\x0B"
    "\x0C"
    "\x0E-\x1F"
    "]"
)


def clean_text(value: Any) -> str:
    if value is None:
        return ""

    text = str(value)
    text = INVALID_XML_CHARACTERS.sub("", text)
    text = text.replace("\xa0", " ")

    return text.strip()


def normalize_label(value: Any) -> str:
    text = clean_text(value)
    text = re.sub(r"\s+", " ", text)

    return text.rstrip(":").strip().casefold()


def create_short_key(
    request_text: str,
    max_words: int = 5,
) -> str:
    text = clean_text(request_text)

    if not text:
        return "YENI_TALEP"

    # "Talep Konusu:" gibi ilk satır etiketlerini kaldırır.
    text = re.sub(
        r"(?im)^\s*talep\s+konusu\s*:\s*",
        "",
        text,
    )

    text = text.translate(TURKISH_CHARACTER_MAP)
    text = unicodedata.normalize("NFKD", text)
    text = text.encode("ascii", "ignore").decode("ascii")

    words = re.findall(r"[A-Za-z0-9]+", text.upper())

    ignored_words = {
        "VE",
        "ILE",
        "ICIN",
        "BIR",
        "BU",
        "SU",
        "OLARAK",
        "YAPILAN",
        "YAPILMASI",
        "SAGLANMASI",
        "DUZELTILMESI",
        "GEREKMEKTEDIR",
        "ISTENMEKTEDIR",
        "TALEP",
        "KONUSU",
        "SORUNUNUN",
    }

    meaningful_words = [
        word
        for word in words
        if word not in ignored_words
    ]

    selected_words = meaningful_words[:max_words]

    if not selected_words:
        selected_words = words[:max_words]

    return "_".join(selected_words) or "YENI_TALEP"


def create_output_path(
    itsm_no: str,
    short_key: str,
) -> Path:
    DOCUMENTS_DIRECTORY.mkdir(
        parents=True,
        exist_ok=True,
    )

    document_number = clean_text(itsm_no) or "XXX"

    base_name = (
        f"DEVCR_{document_number}_{short_key}"
    )

    output_path = (
        DOCUMENTS_DIRECTORY
        / f"{base_name}.docx"
    )

    counter = 2

    while output_path.exists():
        output_path = (
            DOCUMENTS_DIRECTORY
            / f"{base_name}_{counter}.docx"
        )
        counter += 1

    return output_path


def format_request_date(value: Any) -> str:
    if value is None or not str(value).strip():
        return date.today().strftime("%d.%m.%Y")

    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")

    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")

    raw_value = str(value).strip()

    supported_formats = (
        "%d.%m.%Y",
        "%d/%m/%Y",
        "%Y-%m-%d",
        "%m/%d/%Y",
    )

    for date_format in supported_formats:
        try:
            parsed_date = datetime.strptime(
                raw_value,
                date_format,
            )

            return parsed_date.strftime("%d.%m.%Y")

        except ValueError:
            continue

    raise ValueError(
        f"Desteklenmeyen tarih biçimi: {raw_value}"
    )


def clear_cell(cell: _Cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph_element = paragraph._element

        for child in list(paragraph_element):
            paragraph_element.remove(child)

    if not cell.paragraphs:
        cell.add_paragraph()


def write_cell(
    cell: _Cell,
    value: Any,
) -> None:
    text = clean_text(value)

    clear_cell(cell)

    if not text:
        return

    lines = text.splitlines()
    first_paragraph = cell.paragraphs[0]

    for index, line in enumerate(lines):
        if index == 0:
            paragraph = first_paragraph
        else:
            paragraph = cell.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)

        run = paragraph.add_run(
            clean_text(line)
        )

        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(
            0,
            0,
            255,
        )
        run.italic = True


def find_value_cell(
    document: DocumentType,
    field_label: str,
) -> _Cell:
    expected_label = normalize_label(field_label)

    for table in document.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue

            current_label = normalize_label(
                row.cells[0].text
            )

            if current_label == expected_label:
                return row.cells[1]

    raise KeyError(
        f"Şablonda alan bulunamadı: {field_label}"
    )


def set_request_category(
    document: DocumentType,
    selected_category: str,
) -> None:
    """
    İstek Kategorisi satırını bulur.

    İlk sürümde kategori seçimini Word üretimini engellemeyecek
    şekilde işler. Satır bulunamazsa uyarı verir ve devam eder.
    """

    category = clean_text(selected_category).casefold()

    valid_categories = {
        "standart",
        "normal",
        "acil",
    }

    if category not in valid_categories:
        raise ValueError(
            "Öncelik Standart, Normal veya Acil olmalıdır."
        )

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            cell_texts = [
                normalize_label(cell.text)
                for cell in row.cells
            ]

            row_text = " | ".join(cell_texts)

            if "istek kategorisi" not in row_text:
                continue

            print(
                f"[OK] İstek Kategorisi satırı bulundu: "
                f"tablo={table_index}, satır={row_index}"
            )

            category_names = (
                "Standart",
                "Normal",
                "Acil",
            )

            found_categories = 0
            processed_cells: set[int] = set()

            for cell in row.cells:
                cell_id = id(cell._tc)

                if cell_id in processed_cells:
                    continue

                processed_cells.add(cell_id)

                original_text = clean_text(cell.text)
                normalized_text = normalize_label(original_text)

                for category_name in category_names:
                    if category_name.casefold() not in normalized_text:
                        continue

                    is_selected = (
                        category_name.casefold() == category
                    )

                    prefix = "☒" if is_selected else "☐"

                    write_cell(
                        cell,
                        f"{prefix} {category_name}",
                    )

                    found_categories += 1
                    break

            if found_categories > 0:
                print(
                    f"[OK] İstek Kategorisi seçildi: "
                    f"{selected_category}"
                )
            else:
                print(
                    "[UYARI] Kategori satırı bulundu ancak "
                    "seçenek hücreleri değiştirilemedi."
                )

            return

    print(
        "[UYARI] Şablonda İstek Kategorisi satırı bulunamadı. "
        "Doküman kategori işaretlenmeden oluşturulacak."
    )

def generate_request_document(
    request_data: dict[str, Any],
) -> Path:
    """
    Sabit şirket şablonunu doldurur.

    Şablonun aslı değiştirilmez.
    Çıktı documents klasörüne kaydedilir.
    """

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            "Şirket Word şablonu bulunamadı:\n"
            f"{TEMPLATE_PATH}"
        )

    request_text = clean_text(
        request_data.get("request")
    )

    if not request_text:
        raise ValueError(
            "Talep bölümü boş bırakılamaz."
        )

    prepared_data = {
        **request_data,
        "request_date": format_request_date(
            request_data.get("request_date")
        ),
    }

    short_key = create_short_key(
        request_text
    )

    output_path = create_output_path(
        itsm_no=clean_text(
            prepared_data.get("itsm_no")
        ),
        short_key=short_key,
    )

    document = Document(
        str(TEMPLATE_PATH)
    )

    for data_key, document_label in FIELD_LABELS.items():
        value = prepared_data.get(
            data_key,
            "",
        )

        value_cell = find_value_cell(
            document=document,
            field_label=document_label,
        )

        write_cell(
            value_cell,
            value,
        )

    set_request_category(
        document=document,
        selected_category=prepared_data.get(
            "priority",
            "Standart",
        ),
    )

    document.save(
        str(output_path)
    )

    return output_path